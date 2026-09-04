from __future__ import annotations

import re
import tempfile
import unittest
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from report_merger.core import ReportModel, parse_report
from report_merger.exporters import export_docx, export_xlsx


ROOT = Path(__file__).resolve().parents[1]
SAMPLES = [
    (ROOT / "samples" / "Yen Thuy.docx", "Xã Yên Thủy"),
    (ROOT / "samples" / "Muong Vang.docx", "Xã Mường Vang"),
    (ROOT / "samples" / "Dai Dong.docx", "Xã Đại Đồng"),
    (ROOT / "samples" / "Ngoc Son.docx", "Xã Ngọc Sơn"),
]


def compact(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


class PipelineTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.model = ReportModel()
        cls.model.analyze([(str(path), name) for path, name in SAMPLES])

    def test_structure_matches_samples(self):
        self.assertEqual(len(self.model.sources), 4)
        self.assertEqual(len(self.model.roots()), 5)
        self.assertEqual([len(group.occurrences) for group in self.model.roots()], [4, 4, 4, 4, 4])
        for group in self.model.groups:
            source_ids = [occ.source_id for occ in group.occurrences]
            self.assertEqual(len(source_ids), len(set(source_ids)))

    def test_signature_tables_are_not_content(self):
        table_blocks = [
            block for source in self.model.sources for occ in source.occurrences
            for block in occ.blocks if block.kind == "table"
        ]
        self.assertEqual(len(table_blocks), 3)
        self.assertFalse(any("Nơi nhận" in block.text for block in table_blocks))

    def test_docx_without_numbering_part_is_still_analyzed(self):
        with tempfile.TemporaryDirectory() as folder:
            source_path = Path(folder) / "source.docx"
            no_numbering_path = Path(folder) / "without-numbering.docx"

            document = Document()
            document.add_paragraph("I. KHÁI QUÁT", style="Heading 1")
            document.add_paragraph("Nội dung báo cáo vẫn phải được đọc.")
            document.save(source_path)

            with zipfile.ZipFile(source_path) as source, zipfile.ZipFile(
                no_numbering_path, "w", zipfile.ZIP_DEFLATED
            ) as target:
                for item in source.infolist():
                    if item.filename == "word/numbering.xml":
                        continue
                    data = source.read(item.filename)
                    if item.filename == "word/_rels/document.xml.rels":
                        root = ElementTree.fromstring(data)
                        for relationship in list(root):
                            if relationship.attrib.get("Type", "").endswith("/numbering"):
                                root.remove(relationship)
                        data = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)
                    elif item.filename == "[Content_Types].xml":
                        root = ElementTree.fromstring(data)
                        for override in list(root):
                            if override.attrib.get("PartName") == "/word/numbering.xml":
                                root.remove(override)
                        data = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)
                    target.writestr(item, data)

            report = parse_report(no_numbering_path, 0)

            self.assertEqual(len(report.occurrences), 1)
            self.assertEqual(report.occurrences[0].clean_title, "KHÁI QUÁT")
            self.assertIn("Nội dung báo cáo", report.occurrences[0].content_text)

    def test_export_round_trip_contains_all_body_text(self):
        with tempfile.TemporaryDirectory() as folder:
            docx_path = Path(folder) / "tong_hop.docx"
            xlsx_path = Path(folder) / "tong_hop.xlsx"
            export_docx(self.model, docx_path, "BÁO CÁO TỔNG HỢP")
            export_xlsx(self.model, xlsx_path)
            self.assertGreater(docx_path.stat().st_size, 30_000)
            self.assertGreater(xlsx_path.stat().st_size, 10_000)
            document = Document(docx_path)
            output_parts = []
            for item in document.iter_inner_content():
                if isinstance(item, Paragraph):
                    output_parts.append(item.text)
                elif isinstance(item, Table):
                    output_parts.extend(cell.text for row in item.rows for cell in row.cells)
            output_text = compact("\n".join(output_parts))
            for source in self.model.sources:
                self.assertIn(f"Nguồn: {source.name}", output_text)
                for occ in source.occurrences:
                    for block in occ.blocks:
                        if block.kind == "paragraph" and len(compact(block.text)) >= 8:
                            self.assertIn(compact(block.text), output_text)

    def test_split_parent_distributes_descendants_by_source(self):
        model = ReportModel()
        model.analyze([(str(path), name) for path, name in SAMPLES])
        created = model.split_group(model.roots()[2].id)
        self.assertEqual(len(created), 4)
        for group_id in created:
            group = model.group(group_id)
            self.assertEqual(len(group.occurrences), 1)
            source_id = group.occurrences[0].source_id
            descendants = []
            stack = list(model.children(group.id))
            while stack:
                child = stack.pop()
                descendants.append(child)
                stack.extend(model.children(child.id))
            self.assertTrue(descendants)
            for child in descendants:
                self.assertTrue(all(occ.source_id == source_id for occ in child.occurrences))


if __name__ == "__main__":
    unittest.main()
