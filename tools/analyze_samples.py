from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def para_info(paragraph):
    ppr = paragraph._p.pPr
    outline = None
    num_id = None
    ilvl = None
    if ppr is not None:
        outline_el = ppr.find(qn("w:outlineLvl"))
        if outline_el is not None:
            outline = outline_el.get(qn("w:val"))
        numpr = ppr.find(qn("w:numPr"))
        if numpr is not None:
            num_el = numpr.find(qn("w:numId"))
            lvl_el = numpr.find(qn("w:ilvl"))
            num_id = num_el.get(qn("w:val")) if num_el is not None else None
            ilvl = lvl_el.get(qn("w:val")) if lvl_el is not None else None
    chars = sum(len(r.text) for r in paragraph.runs)
    bold_chars = sum(len(r.text) for r in paragraph.runs if r.bold)
    return {
        "text": paragraph.text.strip(),
        "style": paragraph.style.name if paragraph.style else "",
        "outline": outline,
        "num_id": num_id,
        "ilvl": ilvl,
        "bold_ratio": round(bold_chars / chars, 2) if chars else 0,
        "alignment": str(paragraph.alignment),
    }


def main(paths):
    result = []
    for raw in paths:
        path = Path(raw)
        doc = Document(path)
        rows = []
        for i, p in enumerate(doc.paragraphs):
            info = para_info(p)
            if info["text"]:
                info["index"] = i
                rows.append(info)
        candidates = [
            row for row in rows
            if len(row["text"]) <= 220 and (
                row["bold_ratio"] >= 0.65
                or row["outline"] is not None
                or row["style"].lower().startswith("heading")
                or re.match(r"^(?:[IVXLCDM]+|[A-Z]|\d+(?:\.\d+)*)[.\)]\s+", row["text"], re.I)
            )
        ]
        result.append({
            "file": str(path),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "candidate_headings": candidates,
        })
    json.dump(result, sys.stdout, ensure_ascii=False, indent=2)


if __name__ == "__main__":
    main(sys.argv[1:])
