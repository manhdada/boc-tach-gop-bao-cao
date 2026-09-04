from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from .core import Block, ReportModel, SectionGroup


FONT_NAME = "Times New Roman"


def _set_run_font(run, size=13, bold=None, italic=None, color=None):
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after in (
        ("Title", 16, 0, 12),
        ("Heading 1", 13, 12, 6),
        ("Heading 2", 13, 8, 4),
        ("Heading 3", 13, 6, 3),
        ("Heading 4", 13, 4, 3),
    ):
        style = doc.styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.15
    doc.styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _roman(value: int) -> str:
    pairs = ((1000, "M"), (900, "CM"), (500, "D"), (400, "CD"), (100, "C"),
             (90, "XC"), (50, "L"), (40, "XL"), (10, "X"), (9, "IX"),
             (5, "V"), (4, "IV"), (1, "I"))
    out = []
    for number, symbol in pairs:
        while value >= number:
            out.append(symbol)
            value -= number
    return "".join(out)


def _group_label(model: ReportModel, group: SectionGroup) -> str:
    siblings = [g for g in model.children(group.parent_id) if g.enabled]
    index = siblings.index(group) + 1
    if group.level == 1:
        return f"{_roman(index)}."
    if group.level == 2:
        return f"{index}."
    if group.marker_kind == "compound":
        parent = model.group(group.parent_id) if group.parent_id else None
        parent_siblings = [g for g in model.children(parent.parent_id) if g.enabled] if parent else []
        parent_index = parent_siblings.index(parent) + 1 if parent in parent_siblings else 1
        return f"{parent_index}.{index}."
    if group.level == 3:
        return f"{chr(ord('a') + (index - 1) % 26)})"
    return f"({index})"


def _iter_groups(model: ReportModel, parent_id=None):
    for group in model.children(parent_id):
        if not group.enabled:
            continue
        yield group
        yield from _iter_groups(model, group.id)


def _add_source_label(doc: Document, name: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"Nguồn: {name}")
    _set_run_font(run, 12, bold=True, italic=True, color=(55, 65, 81))


def _copy_paragraph(doc: Document, block: Block):
    paragraph = doc.add_paragraph()
    paragraph.alignment = block.alignment if block.alignment is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    if block.left_indent_pt is not None:
        paragraph.paragraph_format.left_indent = Pt(max(-10, min(block.left_indent_pt, 100)))
    if block.first_line_indent_pt is not None:
        paragraph.paragraph_format.first_line_indent = Pt(max(-30, min(block.first_line_indent_pt, 60)))
    runs = block.runs or []
    if not runs:
        runs = [type("Run", (), {"text": block.text, "bold": None, "italic": None, "underline": None})()]
    for source_run in runs:
        run = paragraph.add_run(source_run.text)
        _set_run_font(run, 13, source_run.bold, source_run.italic)
        if source_run.underline:
            run.underline = True


def _set_table_geometry(table, usable_width_dxa=9354):
    rows = table.findall(qn("w:tr"))
    if not rows:
        return
    first_cells = rows[0].findall(qn("w:tc"))
    col_count = max(1, len(first_cells))
    width = usable_width_dxa // col_count
    tbl_pr = table.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(usable_width_dxa))
    for row in rows:
        for cell in row.findall(qn("w:tc")):
            tc_pr = cell.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = OxmlElement("w:tcPr")
                cell.insert(0, tc_pr)
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def _copy_table(doc: Document, block: Block):
    if block.table_element is None:
        return
    table = copy.deepcopy(block.table_element)
    _set_table_geometry(table)
    body = doc._body._element
    body.insert(max(0, len(body) - 1), table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def export_docx(model: ReportModel, output_path: str | Path, title: str) -> Path:
    output_path = Path(output_path)
    doc = Document()
    _configure_document(doc)
    title_paragraph = doc.add_paragraph(style="Title")
    title_run = title_paragraph.add_run(title.strip() or "BÁO CÁO TỔNG HỢP")
    _set_run_font(title_run, 16, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    subrun = subtitle.add_run(f"Tổng hợp nguyên văn từ {len(model.sources)} báo cáo")
    _set_run_font(subrun, 12, italic=True, color=(89, 89, 89))

    for group in _iter_groups(model):
        label = _group_label(model, group)
        heading = doc.add_paragraph(style=f"Heading {min(group.level, 4)}")
        heading_run = heading.add_run(f"{label} {group.title}")
        _set_run_font(heading_run, 13, bold=True)
        occurrences = sorted(group.occurrences, key=lambda item: item.source_order)
        content_occurrences = [occ for occ in occurrences if occ.blocks]
        for occ in content_occurrences:
            _add_source_label(doc, occ.source_name)
            for block in occ.blocks:
                if block.kind == "table":
                    _copy_table(doc, block)
                else:
                    _copy_paragraph(doc, block)
    doc.save(output_path)
    return output_path


def export_xlsx(model: ReportModel, output_path: str | Path) -> Path:
    output_path = Path(output_path)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Nội dung tổng hợp"
    headers = ["STT", "Cấp", "Đường dẫn mục", "Tên mục", "Nguồn", "Nội dung nguyên văn"]
    sheet.append(headers)
    for cell in sheet[1]:
        cell.font = Font(name=FONT_NAME, bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    row_number = 1

    def walk(parent_id=None, parents=None):
        nonlocal row_number
        parents = parents or []
        for group in model.children(parent_id):
            if not group.enabled:
                continue
            path = parents + [group.title]
            for occ in sorted(group.occurrences, key=lambda item: item.source_order):
                if not occ.blocks:
                    continue
                sheet.append([row_number, group.level, " > ".join(path), group.title,
                              occ.source_name, occ.content_text])
                row_number += 1
            walk(group.id, path)

    walk()
    widths = [8, 8, 48, 38, 30, 90]
    for idx, width in enumerate(widths, 1):
        sheet.column_dimensions[get_column_letter(idx)].width = width
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.font = Font(name=FONT_NAME, size=12)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    workbook.save(output_path)
    return output_path
